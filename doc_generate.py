import json
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def create_assignment_docx(filename):
    doc = Document()

    # --- HELPER FUNCTIONS ---
    def set_font(run, font_name='Calibri', font_size=11, bold=False, color=None):
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        if color:
            run.font.color.rgb = color

    def add_code_block(doc, code_text):
        """Adds a paragraph styled like a code block"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.2)

        # Split by newlines to keep formatting
        lines = code_text.split('\n')
        for line in lines:
            run = p.add_run(line + '\n')
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0, 51, 102)  # Dark Blue for code

    def add_table_with_headers(doc, headers, data):
        """Adds a table with bold headers and 'Table Grid' style"""
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        # Header Row
        hdr_cells = table.rows[0].cells
        for i, header_text in enumerate(headers):
            p = hdr_cells[i].paragraphs[0]
            run = p.add_run(header_text)
            run.bold = True

        # Data Rows
        for row_data in data:
            row_cells = table.add_row().cells
            for i, cell_text in enumerate(row_data):
                row_cells[i].text = str(cell_text)

    # --- DOCUMENT CONTENT ---

    # 1. Title
    title = doc.add_heading('Assignment: Warehouse & Package Management API', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2. Objective
    doc.add_heading('🎯 Objective', level=1)
    p_obj = doc.add_paragraph()
    run_obj = p_obj.add_run(
        "Goal: Build a robust RESTful API using FastAPI to manage a logistics system where Packages belong to specific Warehouses.")

    doc.add_paragraph(
        "This assignment moves beyond simple CRUD operations. It challenges you to implement business logic that spans multiple entities.",
        style='Normal')

    doc.add_heading('🧠 Key Concepts You Will Practice', level=2)
    concepts = [
        "Foreign Key (FK) Relationships: Linking data models (One-to-Many).",
        "Denormalization: Storing copies of parent data in child records for faster reads.",
        "Derived Values: Calculating fields (shipping_cost) dynamically.",
        "Event-Driven Updates: Automatically updating child records when a parent record changes.",
        "Cascading Deletes: Cleaning up related data automatically."
    ]
    for concept in concepts:
        doc.add_paragraph(concept, style='List Bullet')

    # 3. Data Entities
    doc.add_heading('📦 Data Entities', level=1)
    doc.add_paragraph("You need to design two database models (SQLAlchemy) and their corresponding Pydantic schemas.")

    doc.add_heading('1. The Warehouse (Parent)', level=2)
    warehouse_headers = ['Field', 'Type', 'Constraints', 'Description']
    warehouse_data = [
        ['id', 'Integer', 'PK, Auto-increment', 'Unique identifier.'],
        ['name', 'String', 'Unique', 'The name of the facility.'],
        ['city', 'String', 'None', 'Location of the warehouse.'],
        ['handling_fee', 'Float', '0.0 - 20.0', 'Percentage fee added to shipping.']
    ]
    add_table_with_headers(doc, warehouse_headers, warehouse_data)

    doc.add_heading('2. The Package (Child)', level=2)
    package_headers = ['Field', 'Type', 'Constraints', 'Description']
    package_data = [
        ['id', 'Integer', 'PK, Auto-increment', 'Unique identifier.'],
        ['description', 'String', 'None', 'Content description.'],
        ['weight', 'Float', '> 0', 'Weight in kg.'],
        ['shipping_cost', 'Float', 'Calculated', 'Derived value. Read-only.'],
        ['warehouse_id', 'Integer', 'FK', 'Links to Warehouse ID.'],
        ['warehouse_name', 'String', 'Copied', 'Denormalized from Warehouse.'],
        ['handling_fee', 'Float', 'Copied', 'Denormalized from Warehouse.']
    ]
    add_table_with_headers(doc, package_headers, package_data)
    doc.add_paragraph("Relationship: One Warehouse → Many Packages.", style='Intense Quote')

    # 4. API Specification
    doc.add_heading('🛠 API Specification', level=1)

    doc.add_heading('A. Warehouse Endpoints', level=2)
    wh_api_headers = ['Method', 'Endpoint', 'Description']
    wh_api_data = [
        ['POST', '/warehouses/', 'Create a new warehouse.'],
        ['GET', '/warehouses/', 'List all warehouses.'],
        ['PUT', '/warehouses/{id}', 'Update attributes. Triggers logic.'],
        ['DELETE', '/warehouses/{id}', 'Delete warehouse. Triggers cascade.']
    ]
    add_table_with_headers(doc, wh_api_headers, wh_api_data)

    doc.add_heading('B. Package Endpoints', level=2)
    pkg_api_headers = ['Method', 'Endpoint', 'Description']
    pkg_api_data = [
        ['POST', '/packages/', 'Create package. Must provide warehouse_id.'],
        ['GET', '/packages/', 'List all packages.'],
        ['PUT', '/packages/{id}', 'Update desc/weight. Triggers recalculation.'],
        ['DELETE', '/packages/{id}', 'Delete a package.']
    ]
    add_table_with_headers(doc, pkg_api_headers, pkg_api_data)

    # 5. Business Logic
    doc.add_heading('⚙️ Business Logic & Rules', level=1)

    logic_points = [
        ("Smart Creation", "Copy warehouse.name and handling_fee to the package. Calculate shipping_cost immediately."),
        ("The Cost Formula", "Base Cost = weight * 50\nShipping Cost = Base Cost + (Base Cost * (handling_fee / 100))"),
        ("Synchronization",
         "If a Warehouse is updated (name/fee), find ALL related packages and update their warehouse_name and handling_fee. Recalculate shipping_cost."),
        ("Safety & Validation",
         "Cascade Delete: Deleting a warehouse deletes its packages.\nValidation: Fee 0-20, Weight > 0.\nIntegrity: Cannot create package for non-existent warehouse.")
    ]

    for title, desc in logic_points:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)

    # 6. Examples
    doc.add_heading('💡 Examples & Expected Behavior', level=1)

    doc.add_heading('Scenario 1: Creating Data', level=3)
    doc.add_paragraph("Request: Create a Warehouse")
    add_code_block(doc,
                   'POST /warehouses/\n{\n  "name": "Central Hub",\n  "city": "New York",\n  "handling_fee": 5.0\n}')

    doc.add_paragraph("Request: Create a Package (Linked to ID 1)")
    add_code_block(doc,
                   'POST /packages/\n{\n  "description": "Gaming Laptop",\n  "weight": 2.0,\n  "warehouse_id": 1\n}')

    doc.add_paragraph("Response (Auto-calculated):")
    add_code_block(doc,
                   '{\n  "id": 1,\n  "description": "Gaming Laptop",\n  "weight": 2.0,\n  "shipping_cost": 105.0,\n  "warehouse_id": 1,\n  "warehouse_name": "Central Hub",\n  "handling_fee": 5.0\n}')

    doc.add_heading('Scenario 2: Updating the Parent', level=3)
    doc.add_paragraph("Request: Update Warehouse Fee to 10.0%")
    add_code_block(doc, 'PUT /warehouses/1\n{\n  "handling_fee": 10.0\n}')

    doc.add_paragraph("Effect on Package (when fetched again):")
    add_code_block(doc,
                   'GET /packages/1\n{\n  "description": "Gaming Laptop",\n  "handling_fee": 10.0,\n  "shipping_cost": 110.0\n}')

    # 7. Bonus
    doc.add_heading('⭐ Bonus Challenges', level=1)
    bonuses = [
        "Heavy Items Filter: GET /packages/heavy (returns items > 20kg)",
        "Warehouse Inventory: GET /warehouses/{id}/packages (list items in specific warehouse)"
    ]
    for b in bonuses:
        doc.add_paragraph(b, style='List Bullet')

    # Save
    doc.save(filename)
    return filename


# Generate the file in the project directory
if __name__ == '__main__':
    project_dir = os.path.abspath(os.path.dirname(__file__))
    output_path = os.path.join(project_dir, 'warehouse_assignment.docx')
    print(f"Generating: {output_path}")
    create_assignment_docx(output_path)
